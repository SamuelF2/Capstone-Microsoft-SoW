"""
Finalize router  —  /api/finalize/...

Handles Phase 4: document generation, handoff package creation, and SoW locking.

Endpoints
---------
  POST /api/finalize/{sow_id}/generate-document  Generate DOCX/PDF from SoW content
  GET  /api/finalize/{sow_id}/download            Serve the generated document file
  POST /api/finalize/{sow_id}/handoff             Create / replace the handoff package
  GET  /api/finalize/{sow_id}/handoff             Retrieve the handoff package
  POST /api/finalize/{sow_id}/lock                Finalize and lock the SoW
"""

from __future__ import annotations

import json
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from typing import Any

import database
from auth import CurrentUser
from config import UPLOAD_DIR
from fastapi import APIRouter, HTTPException, Query, status
from fastapi.responses import StreamingResponse
from models import (
    DocumentGenerationResponse,
    HandoffPackagePayload,
    HandoffPackageResponse,
)
from utils.db_helpers import insert_history, require_collaborator

router = APIRouter(prefix="/api/finalize", tags=["finalize"])

# ── Internal helpers ──────────────────────────────────────────────────────────


def _safe_json(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, dict | list):
        return value
    try:
        return json.loads(value)
    except (TypeError, ValueError):
        return value


def _bullet_list(items: Any) -> list[str]:
    """Convert a JSONB list / string into displayable bullet strings."""
    if not items:
        return []
    if isinstance(items, str):
        return [items]
    if isinstance(items, list):
        result = []
        for item in items:
            if isinstance(item, str):
                result.append(item)
            elif isinstance(item, dict):
                label = (
                    item.get("name")
                    or item.get("title")
                    or item.get("description")
                    or json.dumps(item)
                )
                result.append(str(label))
            else:
                result.append(str(item))
        return result
    return [str(items)]


def _generated_dir(sow_id: int) -> Path:
    d = Path(UPLOAD_DIR) / "generated" / str(sow_id)
    d.mkdir(parents=True, exist_ok=True)
    return d


# ── DOCX generation ───────────────────────────────────────────────────────────


def _build_docx(sow: dict, content: dict, review_results: list) -> bytes:
    """Build a DOCX document from structured SoW data. Returns raw bytes."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as err:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="python-docx is not installed. Run: pip install python-docx",
        ) from err

    doc = Document()

    # ── Cover ──────────────────────────────────────────────────────────────────
    title_para = doc.add_heading("STATEMENT OF WORK", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    for i, (lbl, val) in enumerate(
        [
            ("Title", sow.get("title") or ""),
            ("Customer", sow.get("customer_name") or ""),
            ("Methodology", sow.get("methodology") or ""),
            ("Generated", datetime.now(UTC).strftime("%B %d, %Y")),
            ("Opportunity ID", sow.get("opportunity_id") or ""),
        ]
    ):
        meta.rows[i].cells[0].text = lbl
        meta.rows[i].cells[1].text = str(val)

    doc.add_paragraph()

    # ── 1. Executive Summary ──────────────────────────────────────────────────
    doc.add_heading("1. EXECUTIVE SUMMARY", 1)
    exec_sum = content.get("executiveSummary") or content.get("executive_summary") or ""
    if isinstance(exec_sum, dict):
        exec_sum = json.dumps(exec_sum, indent=2)
    doc.add_paragraph(str(exec_sum) if exec_sum else "(No executive summary provided)")

    # ── 2. Scope of Work ─────────────────────────────────────────────────────
    doc.add_heading("2. SCOPE OF WORK", 1)
    scope_data = (
        content.get("scope")
        or content.get("projectScope")
        or content.get("cloudAdoptionScope")
        or {}
    )
    if isinstance(scope_data, list):
        scope_data = {}

    doc.add_heading("2.1 In Scope", 2)
    in_scope = scope_data.get("in_scope") or [] if isinstance(scope_data, dict) else []
    for item in _bullet_list(in_scope):
        doc.add_paragraph(f"• {item}")
    if not in_scope:
        doc.add_paragraph("(Not specified)")

    doc.add_heading("2.2 Out of Scope", 2)
    out_scope = scope_data.get("out_scope") or [] if isinstance(scope_data, dict) else []
    for item in _bullet_list(out_scope):
        doc.add_paragraph(f"• {item}")
    if not out_scope:
        doc.add_paragraph("(Not specified)")

    # ── 3. Approach / Methodology ─────────────────────────────────────────────
    doc.add_heading("3. APPROACH / METHODOLOGY", 1)
    approach = None
    for key in (
        "agileApproach",
        "sureStepMethodology",
        "waterfallApproach",
        "migrationStrategy",
        "workloadAssessment",
    ):
        if content.get(key):
            approach = content[key]
            break
    if approach:
        if isinstance(approach, dict):
            approach = json.dumps(approach, indent=2)
        doc.add_paragraph(str(approach))
    else:
        methodology = sow.get("methodology") or ""
        doc.add_paragraph(f"Methodology: {methodology}" if methodology else "(No approach defined)")

    # ── 4. Deliverables ───────────────────────────────────────────────────────
    doc.add_heading("4. DELIVERABLES", 1)
    deliverables = content.get("deliverables") or content.get("phasesDeliverables") or []
    if deliverables:
        for item in _bullet_list(deliverables):
            doc.add_paragraph(f"• {item}")
    else:
        doc.add_paragraph("(No deliverables defined)")

    # ── 5. Team Structure ─────────────────────────────────────────────────────
    doc.add_heading("5. TEAM STRUCTURE", 1)
    resources = content.get("teamStructure") or content.get("resources") or []
    if resources:
        for item in _bullet_list(resources):
            doc.add_paragraph(f"• {item}")
    else:
        doc.add_paragraph("(No team structure defined)")

    # ── 6. Timeline & Milestones ──────────────────────────────────────────────
    doc.add_heading("6. TIMELINE & MILESTONES", 1)
    timeline = content.get("phasesMilestones") or content.get("timeline") or {}
    if timeline:
        if isinstance(timeline, str):
            doc.add_paragraph(timeline)
        elif isinstance(timeline, dict):
            for k, v in timeline.items():
                doc.add_heading(k.replace("_", " ").title(), 3)
                for item in _bullet_list(v if isinstance(v, list) else [v]):
                    doc.add_paragraph(f"• {item}")
        elif isinstance(timeline, list):
            for item in _bullet_list(timeline):
                doc.add_paragraph(f"• {item}")
    else:
        doc.add_paragraph("(No timeline defined)")

    # ── 7. Pricing ────────────────────────────────────────────────────────────
    doc.add_heading("7. PRICING", 1)
    pricing = content.get("pricing") or {}
    if isinstance(pricing, list) and pricing:
        pricing = pricing[0]
    if pricing and isinstance(pricing, dict):
        total = pricing.get("total_value") or pricing.get("totalValue") or sow.get("deal_value")
        if total:
            try:
                doc.add_paragraph(f"Total Value: ${float(total):,.2f}")
            except (TypeError, ValueError):
                doc.add_paragraph(f"Total Value: {total}")
        breakdown = pricing.get("breakdown") or []
        for item in _bullet_list(breakdown):
            doc.add_paragraph(f"• {item}")
    elif sow.get("deal_value"):
        try:
            doc.add_paragraph(f"Total Value: ${float(sow['deal_value']):,.2f}")
        except (TypeError, ValueError):
            doc.add_paragraph(f"Total Value: {sow['deal_value']}")
    else:
        doc.add_paragraph("(Pricing not defined)")

    # ── 8. Assumptions & Risks ────────────────────────────────────────────────
    doc.add_heading("8. ASSUMPTIONS & RISKS", 1)
    doc.add_heading("8.1 Assumptions", 2)
    assumptions = content.get("assumptions") or []
    if isinstance(assumptions, dict):
        assumptions = assumptions.get("items") or []
    for item in _bullet_list(assumptions if isinstance(assumptions, list) else [assumptions]):
        doc.add_paragraph(f"• {item}")
    if not assumptions:
        doc.add_paragraph("(No assumptions listed)")

    doc.add_heading("8.2 Risks", 2)
    risks = content.get("risks") or []
    if risks:
        for r in risks if isinstance(risks, list) else [risks]:
            if isinstance(r, str):
                doc.add_paragraph(f"• {r}")
            elif isinstance(r, dict):
                severity = r.get("severity") or r.get("level") or ""
                desc = r.get("description") or r.get("risk") or str(r)
                mitigation = r.get("mitigation") or ""
                prefix = f"[{severity.upper()}] " if severity else ""
                doc.add_paragraph(f"• {prefix}{desc}")
                if mitigation:
                    doc.add_paragraph(f"  Mitigation: {mitigation}")
    else:
        doc.add_paragraph("(No risks identified)")

    # ── 9. Customer Responsibilities ──────────────────────────────────────────
    doc.add_heading("9. CUSTOMER RESPONSIBILITIES", 1)
    cust_resp = (
        (scope_data.get("customer_responsibilities") if isinstance(scope_data, dict) else [])
        or content.get("customerResponsibilities")
        or []
    )
    for item in _bullet_list(cust_resp):
        doc.add_paragraph(f"• {item}")
    if not cust_resp:
        doc.add_paragraph("(No customer responsibilities listed)")

    # ── 10. Support & Transition ──────────────────────────────────────────────
    doc.add_heading("10. SUPPORT & TRANSITION", 1)
    support = (
        content.get("supportTransition")
        or content.get("supportHypercare")
        or content.get("supportOperations")
        or ""
    )
    if isinstance(support, dict):
        support = json.dumps(support, indent=2)
    doc.add_paragraph(str(support) if support else "(No support plan defined)")

    # ── Appendix A: Review Conditions ─────────────────────────────────────────
    all_conditions: list[str] = []
    for r in review_results:
        cond = _safe_json(r.get("conditions"))
        if isinstance(cond, list):
            all_conditions.extend([str(c) for c in cond if c])
        elif isinstance(cond, str) and cond:
            all_conditions.append(cond)

    if all_conditions:
        doc.add_heading("APPENDIX A: REVIEW CONDITIONS", 1)
        for cond in all_conditions:
            doc.add_paragraph(f"• {cond}")

    # ── Appendix B: Approval Chain ────────────────────────────────────────────
    if review_results:
        doc.add_heading("APPENDIX B: APPROVAL CHAIN", 1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        for i, h in enumerate(["Reviewer", "Role", "Decision", "Date"]):
            tbl.rows[0].cells[i].text = h
        for r in review_results:
            row_cells = tbl.add_row().cells
            row_cells[0].text = str(r.get("reviewer") or "")
            row_cells[1].text = str(r.get("reviewer_role") or r.get("review_stage") or "")
            row_cells[2].text = str(r.get("decision") or "")
            date_val = r.get("reviewed_at")
            if isinstance(date_val, datetime):
                date_str = date_val.strftime("%b %d, %Y")
            elif date_val:
                date_str = str(date_val)[:10]
            else:
                date_str = ""
            row_cells[3].text = date_str

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── POST /api/finalize/{sow_id}/generate-document ────────────────────────────


@router.post(
    "/{sow_id}/generate-document",
    response_model=DocumentGenerationResponse,
    summary="Generate a DOCX or PDF document from the SoW",
)
async def generate_document(
    sow_id: int,
    current_user: CurrentUser,
    fmt: str = Query(default="docx", alias="format", description="Output format: docx or pdf"),
) -> DocumentGenerationResponse:
    """Generate a document from the approved SoW's structured content.

    Stores the file in ``UPLOAD_DIR/generated/{sow_id}/`` and updates
    the handoff package ``document_path`` if one exists.
    """
    if fmt not in ("docx", "pdf"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="format must be 'docx' or 'pdf'"
        )

    async with database.pg_pool.acquire() as conn:
        await require_collaborator(conn, sow_id=sow_id, user_id=current_user.id)
        sow = await conn.fetchrow("SELECT * FROM sow_documents WHERE id = $1", sow_id)
        if not sow:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="SoW not found")
        if sow["status"] not in ("approved", "finalized"):
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=f"SoW must be 'approved' or 'finalized' to generate a document; currently '{sow['status']}'",
            )

        review_rows = await conn.fetch(
            """
            SELECT rr.reviewer, rr.decision, rr.conditions, rr.reviewed_at, ra.reviewer_role, rr.review_stage
            FROM   review_results rr
            LEFT JOIN review_assignments ra ON ra.sow_id = rr.sow_id AND ra.user_id = rr.reviewer_user_id
            WHERE  rr.sow_id = $1
            ORDER  BY rr.reviewed_at
            """,
            sow_id,
        )

    content = _safe_json(sow["content"]) or {}
    review_results = [dict(r) for r in review_rows]

    # ── Build DOCX ────────────────────────────────────────────────────────────
    docx_bytes = _build_docx(dict(sow), content, review_results)

    # ── Persist file ──────────────────────────────────────────────────────────
    safe_title = (
        "".join(c if c.isalnum() or c in "-_ " else "" for c in (sow["title"] or "SoW"))
        .strip()
        .replace(" ", "-")[:60]
    )
    file_name = f"SoW-{safe_title}.docx"
    out_dir = _generated_dir(sow_id)
    file_path = out_dir / file_name

    file_path.write_bytes(docx_bytes)

    # ── Update handoff package document_path if it exists ─────────────────────
    async with database.pg_pool.acquire() as conn:
        await conn.execute(
            "UPDATE handoff_packages SET document_path = $1 WHERE sow_id = $2",
            str(file_path),
            sow_id,
        )

    return DocumentGenerationResponse(
        file_path=str(file_path),
        file_name=file_name,
        format="docx",
        size_bytes=len(docx_bytes),
    )


# ── GET /api/finalize/{sow_id}/download ──────────────────────────────────────


@router.get(
    "/{sow_id}/download",
    summary="Download the generated document",
)
async def download_document(sow_id: int, current_user: CurrentUser):
    """Serve the most recently generated document for a SoW."""
    async with database.pg_pool.acquire() as conn:
        await require_collaborator(conn, sow_id=sow_id, user_id=current_user.id)
    # Find latest file in generated directory
    out_dir = _generated_dir(sow_id)
    candidates = list(out_dir.glob("*.docx")) + list(out_dir.glob("*.pdf"))
    if not candidates:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No generated document found. Generate one first.",
        )
    latest = max(candidates, key=lambda p: p.stat().st_mtime)

    ext = latest.suffix.lower()
    if ext == ".pdf":
        media_type = "application/pdf"
    else:
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    file_bytes = latest.read_bytes()
    headers = {"Content-Disposition": f'attachment; filename="{latest.name}"'}
    return StreamingResponse(BytesIO(file_bytes), media_type=media_type, headers=headers)


# ── POST /api/finalize/{sow_id}/handoff ──────────────────────────────────────


@router.post(
    "/{sow_id}/handoff",
    response_model=HandoffPackageResponse,
    summary="Create or replace the handoff package",
)
async def create_handoff(
    sow_id: int,
    payload: HandoffPackagePayload,
    current_user: CurrentUser,
) -> HandoffPackageResponse:
    """Build and persist the handoff package for the approved SoW.

    Automatically bundles approved scope, deliverables, resource plan,
    risk register, and review decisions alongside the supplied payload.
    """
    async with database.pg_pool.acquire() as conn:
        await require_collaborator(conn, sow_id=sow_id, user_id=current_user.id)
        sow = await conn.fetchrow("SELECT * FROM sow_documents WHERE id = $1", sow_id)
        if not sow:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="SoW not found")
        if sow["status"] not in ("approved", "finalized"):
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="SoW must be 'approved' or 'finalized' to create a handoff package",
            )

        review_rows = await conn.fetch(
            """
            SELECT rr.reviewer, rr.decision, rr.conditions, rr.reviewed_at, ra.reviewer_role
            FROM   review_results rr
            LEFT JOIN review_assignments ra ON ra.sow_id = rr.sow_id AND ra.user_id = rr.reviewer_user_id
            WHERE  rr.sow_id = $1
            ORDER  BY rr.reviewed_at
            """,
            sow_id,
        )

    content = _safe_json(sow["content"]) or {}
    scope_data = content.get("scope") or content.get("projectScope") or {}
    if isinstance(scope_data, list):
        scope_data = {}

    # ── Gather conditions from approved-with-conditions decisions ─────────────
    conditions_to_address: list[str] = []
    review_decisions: list[dict] = []
    for r in review_rows:
        cond = _safe_json(r["conditions"])
        if isinstance(cond, list):
            conditions_to_address.extend([str(c) for c in cond if c])
        elif isinstance(cond, str) and cond:
            conditions_to_address.append(cond)
        date_val = r["reviewed_at"]
        review_decisions.append(
            {
                "reviewer": r["reviewer"],
                "role": r["reviewer_role"] or "",
                "decision": r["decision"],
                "date": date_val.isoformat()
                if isinstance(date_val, datetime)
                else str(date_val or ""),
            }
        )

    # ── Build package_data ────────────────────────────────────────────────────
    package_data: dict[str, Any] = {
        "sow_summary": {
            "title": sow["title"],
            "customer_name": sow["customer_name"],
            "methodology": sow["methodology"],
            "deal_value": float(sow["deal_value"]) if sow["deal_value"] else None,
            "esap_level": sow["esap_level"],
            "opportunity_id": sow["opportunity_id"],
        },
        "approved_scope": {
            "in_scope": scope_data.get("in_scope") or [] if isinstance(scope_data, dict) else [],
            "out_scope": scope_data.get("out_scope") or [] if isinstance(scope_data, dict) else [],
        },
        "deliverables": content.get("deliverables") or [],
        "resource_plan": content.get("teamStructure") or content.get("resources") or [],
        "risk_register": content.get("risks") or [],
        "review_decisions": review_decisions,
        "conditions_to_address": conditions_to_address,
        "timeline": content.get("phasesMilestones") or content.get("timeline") or {},
        "customer_responsibilities": (
            scope_data.get("customer_responsibilities") if isinstance(scope_data, dict) else []
        )
        or [],
        # From payload
        "delivery_team": payload.delivery_team,
        "key_contacts": payload.key_contacts,
        "kickoff_date": payload.kickoff_date,
        "special_instructions": payload.special_instructions,
        "notes": payload.notes,
    }

    # ── Attachment manifest (Phase 4) ────────────────────────────────────
    async with database.pg_pool.acquire() as conn:
        attachment_rows = await conn.fetch(
            """
            SELECT id, original_name, document_type, stage_key, file_size, uploaded_at
            FROM sow_attachments
            WHERE sow_id = $1
            ORDER BY uploaded_at
            """,
            sow_id,
        )
    package_data["attachments"] = [
        {
            "id": a["id"],
            "original_name": a["original_name"],
            "document_type": a["document_type"],
            "stage_key": a["stage_key"],
            "file_size": a["file_size"],
            "uploaded_at": a["uploaded_at"].isoformat() if a["uploaded_at"] else None,
            "download_url": f"/api/attachments/{a['id']}/download",
        }
        for a in attachment_rows
    ]

    # ── Check for existing document path ──────────────────────────────────────
    async with database.pg_pool.acquire() as conn:
        existing_doc_path = await conn.fetchval(
            "SELECT document_path FROM handoff_packages WHERE sow_id = $1 LIMIT 1",
            sow_id,
        )

        # Delete existing packages for this SoW (replace semantics)
        await conn.execute("DELETE FROM handoff_packages WHERE sow_id = $1", sow_id)

        row = await conn.fetchrow(
            """
            INSERT INTO handoff_packages (sow_id, created_by, document_path, package_data)
            VALUES ($1, $2, $3, $4::jsonb)
            RETURNING *
            """,
            sow_id,
            current_user.id,
            existing_doc_path,
            json.dumps(package_data),
        )

    return HandoffPackageResponse(
        id=row["id"],
        sow_id=row["sow_id"],
        created_by=row["created_by"],
        document_path=row["document_path"],
        package_data=_safe_json(row["package_data"]) or {},
        created_at=row["created_at"],
    )


# ── GET /api/finalize/{sow_id}/handoff ───────────────────────────────────────


@router.get(
    "/{sow_id}/handoff",
    response_model=HandoffPackageResponse,
    summary="Retrieve the handoff package for a SoW",
)
async def get_handoff(sow_id: int, current_user: CurrentUser) -> HandoffPackageResponse:
    """Return the most recent handoff package for a SoW."""
    async with database.pg_pool.acquire() as conn:
        await require_collaborator(conn, sow_id=sow_id, user_id=current_user.id)
        row = await conn.fetchrow(
            "SELECT * FROM handoff_packages WHERE sow_id = $1 ORDER BY created_at DESC LIMIT 1",
            sow_id,
        )
    if not row:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="No handoff package found for this SoW"
        )
    return HandoffPackageResponse(
        id=row["id"],
        sow_id=row["sow_id"],
        created_by=row["created_by"],
        document_path=row["document_path"],
        package_data=_safe_json(row["package_data"]) or {},
        created_at=row["created_at"],
    )


# ── POST /api/finalize/{sow_id}/lock ─────────────────────────────────────────


@router.post(
    "/{sow_id}/lock",
    summary="Finalize and lock the SoW",
)
async def lock_sow(sow_id: int, current_user: CurrentUser) -> dict:
    """Transition the SoW to ``finalized`` status and permanently lock it.

    Prerequisites:
    - SoW must be ``approved``
    - A handoff package must exist
    - A generated document must exist on disk

    After locking, PATCH and DELETE are rejected by the sow router guards.
    """
    async with database.pg_pool.acquire() as conn:
        await require_collaborator(conn, sow_id=sow_id, user_id=current_user.id)
        sow = await conn.fetchrow("SELECT * FROM sow_documents WHERE id = $1", sow_id)
        if not sow:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="SoW not found")
        if sow["status"] != "approved":
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=f"SoW must be 'approved' to lock; currently '{sow['status']}'",
            )

        # Verify handoff package exists
        handoff = await conn.fetchrow(
            "SELECT id, document_path FROM handoff_packages WHERE sow_id = $1 LIMIT 1",
            sow_id,
        )
        if not handoff:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="A handoff package must be created before locking",
            )

        # Verify generated document exists
        doc_path = handoff["document_path"]
        if not doc_path or not Path(doc_path).exists():
            # Also check the generated dir for any file
            gen_dir = _generated_dir(sow_id)
            candidates = list(gen_dir.glob("*.docx")) + list(gen_dir.glob("*.pdf"))
            if not candidates:
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail="A document must be generated before locking",
                )

    # Gate: block finalization if any COAs are outstanding
    async with database.pg_pool.acquire() as conn:
        outstanding_count = await conn.fetchval(
            """
            SELECT count(*) FROM conditions_of_approval
            WHERE sow_id = $1 AND status NOT IN ('resolved', 'waived')
            """,
            sow_id,
        )
        if outstanding_count and outstanding_count > 0:
            outstanding_rows = await conn.fetch(
                """
                SELECT id, condition_text, status, category, priority
                FROM conditions_of_approval
                WHERE sow_id = $1 AND status NOT IN ('resolved', 'waived')
                ORDER BY priority DESC
                """,
                sow_id,
            )
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail={
                    "message": f"{outstanding_count} condition(s) of approval still outstanding",
                    "outstanding_conditions": [dict(r) for r in outstanding_rows],
                },
            )

    async with database.pg_pool.acquire() as conn, conn.transaction():
        now = datetime.now(UTC)
        await conn.execute(
            """
            UPDATE sow_documents
            SET    status       = 'finalized',
                   finalized_at = $1,
                   finalized_by = $2,
                   updated_at   = $1
            WHERE  id = $3
            """,
            now,
            current_user.id,
            sow_id,
        )
        await conn.execute(
            "UPDATE sow_workflow SET current_stage = 'finalized', updated_at = NOW() WHERE sow_id = $1",
            sow_id,
        )
        await insert_history(
            conn,
            sow_id,
            current_user.id,
            "finalized",
            {
                "finalized_by_email": current_user.email,
            },
        )

    return {
        "finalized": True,
        "sow_id": sow_id,
        "status": "finalized",
        "finalized_at": now.isoformat(),
    }
