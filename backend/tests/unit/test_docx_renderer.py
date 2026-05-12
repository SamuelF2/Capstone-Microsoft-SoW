"""Unit tests for ``services.docx_renderer``.

Exercises ``render_sow_to_docx`` against the four seed content templates
shipped in ``backend/content_templates/`` plus a handful of edge-case
fixtures, asserting the rendered DOCX mirrors the structure of the in-app
SoW viewer (``frontend/components/sow/SoWDocumentReader.js``).
"""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from services.docx_renderer import (
    SECTION_LABELS,
    _humanize_key,
    _is_simple_text_array,
    _is_table_candidate,
    _ordered_section_keys,
    _primary_text,
    render_sow_to_docx,
)

TEMPLATES_DIR = Path(__file__).resolve().parents[2] / "content_templates"


# ── Fixtures ──────────────────────────────────────────────────────────────────


def _load_template(name: str) -> dict[str, Any]:
    """Load a template JSON and return its ``template_data`` payload."""
    with open(TEMPLATES_DIR / f"{name}.json", encoding="utf-8") as fh:
        return json.load(fh)["template_data"]


def _sow_row(title: str = "Test SoW", methodology: str = "Test") -> dict[str, Any]:
    return {
        "id": 1,
        "title": title,
        "customer_name": "Contoso Ltd",
        "methodology": methodology,
        "opportunity_id": "OPP-12345",
        "deal_value": 250000,
        "status": "approved",
    }


def _render_to_doc(
    content: dict[str, Any],
    *,
    sow: dict[str, Any] | None = None,
    review_results: list[dict[str, Any]] | None = None,
) -> Any:
    """Render and re-open the DOCX so we can introspect it."""
    raw = render_sow_to_docx(sow or _sow_row(), content, review_results or [])
    return Document(BytesIO(raw))


def _heading_texts(doc: Any) -> list[str]:
    return [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]


def _level1_texts(doc: Any) -> list[str]:
    return [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]


def _all_text(doc: Any) -> str:
    """Concatenate every paragraph and every cell so we can assert content presence."""
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ── Pure helper tests ─────────────────────────────────────────────────────────


class TestHumanizeKey:
    @pytest.mark.parametrize(
        "raw,expected",
        [
            # camelCase → each uppercase letter starts a new word.
            ("executiveSummary", "Executive Summary"),
            ("inScope", "In Scope"),
            ("outOfScope", "Out Of Scope"),
            ("phasesDeliverables", "Phases Deliverables"),
            # snake_case → underscores become spaces; only the first letter
            # is uppercased (matches the React ``humanizeKey``).
            ("project_scope", "Project scope"),
            ("ABC", "ABC"),
            ("", ""),
        ],
    )
    def test_humanize(self, raw: str, expected: str) -> None:
        assert _humanize_key(raw) == expected


class TestPrimaryText:
    def test_picks_text_first(self) -> None:
        key, val = _primary_text({"id": "a", "text": "hello", "title": "ignored"})
        assert (key, val) == ("text", "hello")

    def test_falls_through_to_story(self) -> None:
        key, val = _primary_text({"id": "a", "story": "as a user…"})
        assert (key, val) == ("story", "as a user…")

    def test_returns_none_when_absent(self) -> None:
        key, val = _primary_text({"id": "a", "foo": "bar"})
        assert (key, val) == (None, None)

    def test_skips_empty_strings(self) -> None:
        key, val = _primary_text({"text": "  ", "title": "Real Title"})
        assert (key, val) == ("title", "Real Title")


class TestTableCandidate:
    def test_uniform_keys_short_text(self) -> None:
        items = [
            {"id": "1", "role": "PM", "name": "Alice", "allocation": "100%"},
            {"id": "2", "role": "Dev", "name": "Bob", "allocation": "50%"},
        ]
        ok, cols = _is_table_candidate(items)
        assert ok is True
        # ``id`` is filtered out.
        assert cols == ["role", "name", "allocation"]

    def test_too_few_items(self) -> None:
        ok, _ = _is_table_candidate([{"a": 1, "b": 2}])
        assert ok is False

    def test_too_many_columns(self) -> None:
        items = [
            {f"col_{i}": str(i) for i in range(7)},
            {f"col_{i}": str(i + 10) for i in range(7)},
        ]
        ok, _ = _is_table_candidate(items)
        assert ok is False

    def test_long_cell_text_disqualifies(self) -> None:
        items = [
            {"epic": "X", "story": "A" * 150, "priority": "high"},
            {"epic": "Y", "story": "B" * 150, "priority": "low"},
        ]
        ok, _ = _is_table_candidate(items)
        assert ok is False

    def test_heterogeneous_keys(self) -> None:
        items = [
            {"a": 1, "b": 2},
            {"a": 1, "c": 3},
        ]
        ok, _ = _is_table_candidate(items)
        assert ok is False


class TestSimpleTextArray:
    def test_id_plus_text_is_simple(self) -> None:
        items = [{"id": "is-1", "text": "Scope item"}, {"id": "is-2", "text": "Another"}]
        assert _is_simple_text_array(items) is True

    def test_extra_field_disqualifies(self) -> None:
        items = [
            {"id": "is-1", "text": "Scope item", "owner": "Microsoft"},
            {"id": "is-2", "text": "Another", "owner": "Customer"},
        ]
        assert _is_simple_text_array(items) is False


class TestOrderedSectionKeys:
    def test_known_keys_use_canonical_order(self) -> None:
        content = {
            "pricing": {"breakdown": []},
            "executiveSummary": {"content": "x"},
            "projectScope": {"inScope": [{"id": "1", "text": "a"}]},
        }
        result = _ordered_section_keys(content)
        assert result == ["executiveSummary", "projectScope", "pricing"]

    def test_unknown_keys_appended(self) -> None:
        content = {
            "customSection": "hello",
            "executiveSummary": {"content": "x"},
        }
        result = _ordered_section_keys(content)
        assert result == ["executiveSummary", "customSection"]

    def test_empty_sections_filtered_out(self) -> None:
        content = {
            "executiveSummary": {"content": "x"},
            "risks": [],
            "deliverables": None,
        }
        result = _ordered_section_keys(content)
        assert result == ["executiveSummary"]

    def test_metadata_keys_excluded(self) -> None:
        content = {
            "sowTitle": "Title",
            "customerName": "Contoso",
            "executiveSummary": {"content": "x"},
        }
        result = _ordered_section_keys(content)
        assert result == ["executiveSummary"]


# ── End-to-end template rendering ─────────────────────────────────────────────


class TestRenderAllTemplates:
    """Each seed template must round-trip into a non-empty DOCX with structured headings."""

    @pytest.mark.parametrize(
        "template_name",
        ["agile-sprint-delivery", "cloud-adoption", "sure-step-365", "waterfall"],
    )
    def test_renders_without_error(self, template_name: str) -> None:
        content = _load_template(template_name)
        doc = _render_to_doc(content)
        # Cover title is present.
        assert any(p.text == "STATEMENT OF WORK" for p in doc.paragraphs)
        # At least one level-1 section heading.
        assert len(_level1_texts(doc)) >= 1

    @pytest.mark.parametrize(
        "template_name",
        ["agile-sprint-delivery", "cloud-adoption", "sure-step-365", "waterfall"],
    )
    def test_every_non_empty_top_level_key_gets_heading(self, template_name: str) -> None:
        content = _load_template(template_name)
        doc = _render_to_doc(content)
        headings = _level1_texts(doc)
        # Every key in canonical order that's present in content should appear
        # as a level-1 heading (either directly or absorbed via nesting; the
        # dedupe logic guarantees no duplicates).
        seen_labels = {h.split(". ", 1)[1] if ". " in h else h for h in headings}
        for k in _ordered_section_keys(content):
            expected = (SECTION_LABELS.get(k) or _humanize_key(k)).upper()
            # The key is either rendered directly OR absorbed by a parent section.
            absorbed = any(
                isinstance(content.get(parent), dict) and k in content[parent]
                for parent in content
                if parent != k
            )
            if not absorbed:
                assert expected in seen_labels, (
                    f"missing heading for {k!r} ({expected}) in {template_name}"
                )

    @pytest.mark.parametrize(
        "template_name",
        ["agile-sprint-delivery", "cloud-adoption", "sure-step-365", "waterfall"],
    )
    def test_no_raw_json_blobs_leak(self, template_name: str) -> None:
        """Bullet items must render their ``text``/``story``/``role`` field — not raw JSON."""
        content = _load_template(template_name)
        doc = _render_to_doc(content)
        all_text = _all_text(doc)
        # No raw JSON dict serialization should appear in body text.  We look
        # for the tell-tale pattern: a substring containing a quoted ``"id":``
        # immediately followed by another quoted key (the old generator's
        # ``json.dumps(item)`` fallback).
        assert '"id":' not in all_text, (
            f"raw JSON leaked into {template_name} document:\n"
            f"{[line for line in all_text.splitlines() if 'id' in line][:5]}"
        )
        # And no literal "{'id'" Python-repr style either.
        assert "{'id'" not in all_text

    @pytest.mark.parametrize(
        "template_name",
        ["agile-sprint-delivery", "cloud-adoption", "sure-step-365", "waterfall"],
    )
    def test_internal_ids_not_used_as_labels(self, template_name: str) -> None:
        """No table column or sub-heading should be labelled ``Id`` / ``_id`` / ``Uid``."""
        content = _load_template(template_name)
        doc = _render_to_doc(content)
        for table in doc.tables:
            for cell in table.rows[0].cells:
                assert cell.text.lower().strip() not in {"id", "_id", "uid"}
        # Sub-section headings:
        for p in doc.paragraphs:
            if p.style.name.startswith("Heading"):
                assert not p.text.lower().strip().endswith(" id"), (
                    f"heading uses ID label: {p.text!r}"
                )

    @pytest.mark.parametrize(
        "template_name",
        ["agile-sprint-delivery", "cloud-adoption", "sure-step-365", "waterfall"],
    )
    def test_in_scope_text_appears(self, template_name: str) -> None:
        """In-scope bullet text must surface — fixing the snake_case key bug."""
        content = _load_template(template_name)
        # Locate first inScope item across the various scope keys.
        first_in_scope_text = None
        for key in ("projectScope", "cloudAdoptionScope", "scope"):
            scope = content.get(key)
            if isinstance(scope, dict):
                items = scope.get("inScope") or scope.get("in_scope") or []
                if items and isinstance(items[0], dict):
                    first_in_scope_text = items[0].get("text")
                    break
        assert first_in_scope_text, f"template {template_name} has no inScope to test"
        doc = _render_to_doc(content)
        all_text = _all_text(doc)
        assert first_in_scope_text in all_text


# ── Targeted behavior tests ──────────────────────────────────────────────────


class TestTablePromotion:
    def test_team_structure_members_is_a_table(self) -> None:
        content = _load_template("agile-sprint-delivery")
        doc = _render_to_doc(content)
        # teamStructure.members is a uniform 4-column dict array → table.
        # Find a table whose header row matches the expected columns.
        member = content["teamStructure"]["members"][0]
        expected_cols = {_humanize_key(k) for k in member if k != "id"}
        found = False
        for table in doc.tables:
            header = {c.text for c in table.rows[0].cells}
            if expected_cols.issubset(header):
                found = True
                break
        assert found, "teamStructure.members did not render as a Word table"

    def test_pricing_breakdown_is_a_table(self) -> None:
        content = _load_template("agile-sprint-delivery")
        doc = _render_to_doc(content)
        breakdown_item = content["pricing"]["breakdown"][0]
        expected_cols = {_humanize_key(k) for k in breakdown_item if k != "id"}
        found = False
        for table in doc.tables:
            header = {c.text for c in table.rows[0].cells}
            if expected_cols.issubset(header):
                found = True
                break
        assert found, "pricing.breakdown did not render as a Word table"

    def test_product_backlog_falls_back_to_cards(self) -> None:
        """Long ``story`` field disqualifies productBacklog from table promotion."""
        content = _load_template("agile-sprint-delivery")
        doc = _render_to_doc(content)
        # No table should have ``Story`` as a header column (the long-text
        # guard kicks in).  But the story text itself should still appear.
        for table in doc.tables:
            headers = {c.text for c in table.rows[0].cells}
            assert "Story" not in headers, "productBacklog wrongly promoted to a table"
        # Spot-check that the first story text surfaces somewhere.
        first_story = content["productBacklog"][0]["story"]
        assert first_story in _all_text(doc)

    def test_deliverables_falls_back_to_cards(self) -> None:
        """deliverables has a long ``description`` → no table promotion."""
        content = _load_template("agile-sprint-delivery")
        doc = _render_to_doc(content)
        for table in doc.tables:
            headers = {c.text for c in table.rows[0].cells}
            assert "Description" not in headers or "Title" not in headers, (
                "deliverables wrongly promoted to a table"
            )
        first_title = content["deliverables"][0]["title"]
        assert first_title in _all_text(doc)


class TestNumbering:
    def test_section_numbering_contiguous_when_filtered(self) -> None:
        """Empty sections must be skipped and numbers stay contiguous."""
        content = {
            "executiveSummary": {"content": "Summary text"},
            "projectScope": None,  # empty — must be skipped
            "deliverables": [],  # empty — must be skipped
            "pricing": {"totalValue": 5000, "breakdown": []},
        }
        doc = _render_to_doc(content)
        level1 = _level1_texts(doc)
        # Should produce "1. EXECUTIVE SUMMARY" and "2. PRICING" — no 3.
        nums = []
        for h in level1:
            try:
                nums.append(int(h.split(".", 1)[0]))
            except (ValueError, IndexError):
                continue
        assert nums == [1, 2], f"section numbers not contiguous: {level1}"

    def test_subsection_numbering(self) -> None:
        content = {
            "projectScope": {
                "inScope": [{"id": "a", "text": "X"}],
                "outOfScope": [{"id": "b", "text": "Y"}],
            },
        }
        doc = _render_to_doc(content)
        level2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
        # Expect "1.1 In Scope" and "1.2 Out Of Scope".
        assert any(t.startswith("1.1 ") for t in level2), level2
        assert any(t.startswith("1.2 ") for t in level2), level2


class TestNestedDedupe:
    def test_risks_nested_in_assumptions_risks_not_duplicated(self) -> None:
        """When ``risks`` is nested inside ``assumptionsRisks``, it should
        not also render as a top-level section."""
        content = {
            "assumptionsRisks": {
                "assumptions": [{"id": "a-1", "text": "Assumption 1"}],
                "risks": [{"id": "r-1", "text": "Risk 1"}],
            },
            "risks": [{"id": "r-2", "text": "Should not appear separately"}],
        }
        doc = _render_to_doc(content)
        level1 = _level1_texts(doc)
        # ``ASSUMPTIONS & RISKS`` should be the only top-level heading for these.
        risks_headings = [h for h in level1 if h.endswith("RISKS")]
        assert len(risks_headings) == 1, f"expected single risks heading, got: {risks_headings}"
        # The text under the absorbed top-level ``risks`` must NOT appear.
        all_text = _all_text(doc)
        assert "Should not appear separately" not in all_text


class TestContentWrapper:
    def test_executive_summary_envelope_unwrapped(self) -> None:
        content = {
            "executiveSummary": {"content": "This is the exec summary body."},
        }
        doc = _render_to_doc(content)
        all_text = _all_text(doc)
        assert "This is the exec summary body." in all_text
        # No spurious "Content:" sub-heading should appear.
        for p in doc.paragraphs:
            if p.style.name.startswith("Heading"):
                assert p.text.strip().lower() != "1.1 content"


class TestCoverMetadata:
    def test_deal_value_formatted_as_currency(self) -> None:
        doc = _render_to_doc({}, sow=_sow_row())
        all_text = _all_text(doc)
        assert "$250,000.00" in all_text

    def test_content_metadata_overrides_sow_row(self) -> None:
        content = {"sowTitle": "Content-Overridden Title"}
        doc = _render_to_doc(content, sow=_sow_row(title="Row Title"))
        all_text = _all_text(doc)
        assert "Content-Overridden Title" in all_text

    def test_handles_empty_deal_value_string(self) -> None:
        """Templates seed deal-value as empty string — must not crash."""
        content = {"dealValue": ""}
        doc = _render_to_doc(content, sow={**_sow_row(), "deal_value": None})
        # No crash and no stray "Deal Value" row when empty.
        for table in doc.tables:
            for row in table.rows:
                labels = [c.text for c in row.cells]
                if labels and labels[0] == "Deal Value":
                    assert labels[1] == ""


class TestAppendices:
    def test_appendix_a_only_when_conditions(self) -> None:
        doc = _render_to_doc({"executiveSummary": "x"}, review_results=[])
        level1 = _level1_texts(doc)
        assert not any("APPENDIX A" in h for h in level1)

    def test_appendix_a_renders_conditions(self) -> None:
        review_results = [
            {
                "reviewer": "Alice",
                "decision": "approved",
                "conditions": ["Update SLA wording", "Confirm pricing"],
                "reviewed_at": None,
                "reviewer_role": "drm",
            },
        ]
        doc = _render_to_doc({"executiveSummary": "x"}, review_results=review_results)
        level1 = _level1_texts(doc)
        assert any("APPENDIX A" in h for h in level1)
        all_text = _all_text(doc)
        assert "Update SLA wording" in all_text
        assert "Confirm pricing" in all_text

    def test_appendix_b_renders_approval_chain(self) -> None:
        review_results = [
            {
                "reviewer": "Bob",
                "decision": "approved",
                "conditions": None,
                "reviewed_at": None,
                "reviewer_role": "delivery_lead",
            },
        ]
        doc = _render_to_doc({"executiveSummary": "x"}, review_results=review_results)
        level1 = _level1_texts(doc)
        assert any("APPENDIX B" in h for h in level1)
        all_text = _all_text(doc)
        assert "Bob" in all_text
        assert "delivery_lead" in all_text


class TestRiskRendering:
    def test_plain_text_risks_render_as_bullets(self) -> None:
        content = {"risks": [{"id": "r-1", "text": "Scope creep"}]}
        doc = _render_to_doc(content)
        all_text = _all_text(doc)
        assert "Scope creep" in all_text
        # No "[SEVERITY]" prefix when severity isn't present.
        assert "[" not in all_text or "Scope creep" not in all_text.split("[")[1]

    def test_severity_and_mitigation_rendered_when_present(self) -> None:
        content = {
            "risks": [
                {
                    "severity": "high",
                    "description": "API instability",
                    "mitigation": "Use mock services",
                },
            ],
        }
        doc = _render_to_doc(content)
        all_text = _all_text(doc)
        assert "[HIGH]" in all_text
        assert "API instability" in all_text
        assert "Mitigation: Use mock services" in all_text


class TestUnknownKeysFallback:
    def test_unknown_section_key_still_rendered(self) -> None:
        content = {
            "executiveSummary": "Summary",
            "customSection": "Some custom content",
        }
        doc = _render_to_doc(content)
        level1 = _level1_texts(doc)
        assert any("CUSTOM SECTION" in h for h in level1)
        assert "Some custom content" in _all_text(doc)
