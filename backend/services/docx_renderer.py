"""DOCX renderer for SoW documents.

Mirrors the structure of the in-app SoW viewer
(``frontend/components/sow/SoWDocumentReader.js``) so the generated Word
document reads the same way reviewers see it on screen: same section order,
same labels, structured tables for tabular arrays, no raw JSON dumps.

Public entry point: :func:`render_sow_to_docx`.

The module is intentionally DB-free — it takes the SoW row, the parsed
``content`` dict, and the list of review results as plain Python values, and
returns DOCX bytes.  This keeps it cheap to unit-test against the four
content templates in ``backend/content_templates/``.
"""

from __future__ import annotations

import json
from datetime import UTC, datetime
from io import BytesIO
from typing import Any

# ── Constants mirrored from the React viewer ──────────────────────────────────
# Keep in sync with frontend/components/sow/SoWDocumentReader.js and
# frontend/components/sow/SoWContentPanel.js.  Small enough that the
# duplication beats indirection through a shared JSON file.

SECTION_ORDER: list[str] = [
    "executiveSummary",
    "projectScope",
    "scope",
    "cloudAdoptionScope",
    "agileApproach",
    "sureStepMethodology",
    "waterfallApproach",
    "migrationStrategy",
    "workloadAssessment",
    "productBacklog",
    "phasesDeliverables",
    "phasesMilestones",
    "deliverables",
    "dataMigration",
    "testingStrategy",
    "teamStructure",
    "supportTransition",
    "supportHypercare",
    "supportOperations",
    "securityCompliance",
    "assumptions",
    "assumptionsRisks",
    "risks",
    "pricing",
]

SECTION_LABELS: dict[str, str] = {
    "executiveSummary": "Executive Summary",
    "projectScope": "Project Scope",
    "scope": "Project Scope",
    "deliverables": "Deliverables",
    "assumptions": "Assumptions",
    "risks": "Risks",
    "pricing": "Pricing",
    "teamStructure": "Team Structure",
    "supportTransition": "Support & Transition",
    "agileApproach": "Agile Approach",
    "productBacklog": "Product Backlog",
    "sureStepMethodology": "Sure Step Methodology",
    "phasesDeliverables": "Phases & Deliverables",
    "dataMigration": "Data Migration",
    "testingStrategy": "Testing Strategy",
    "supportHypercare": "Support & Hypercare",
    "waterfallApproach": "Waterfall Approach",
    "phasesMilestones": "Phases & Milestones",
    "cloudAdoptionScope": "Cloud Adoption Scope",
    "migrationStrategy": "Migration Strategy",
    "workloadAssessment": "Workload Assessment",
    "securityCompliance": "Security & Compliance",
    "supportOperations": "Support & Operations",
    "assumptionsRisks": "Assumptions & Risks",
}

METADATA_FIELDS: tuple[str, ...] = (
    "status",
    "sowTitle",
    "customerName",
    "opportunityId",
    "dealValue",
    "deliveryMethodology",
)

METADATA_LABELS: dict[str, str] = {
    "status": "Status",
    "sowTitle": "Title",
    "customerName": "Customer",
    "opportunityId": "Opportunity ID",
    "dealValue": "Deal Value",
    "deliveryMethodology": "Delivery Methodology",
}

PRIMARY_TEXT_FIELDS: tuple[str, ...] = ("text", "title", "name", "item", "story", "role")

HIDDEN_FIELDS: frozenset[str] = frozenset({"id", "_id", "uid"})

# Table-promotion heuristic guards: keep cells from wrapping into illegible
# narrow columns.
_MAX_TABLE_COLUMNS = 5
_MAX_TABLE_CELL_CHARS = 80


# ── Small helpers ─────────────────────────────────────────────────────────────


def _is_empty(value: Any) -> bool:
    """True for None, empty string, empty list/dict — matches ``isEmpty`` in the viewer."""
    if value is None:
        return True
    if isinstance(value, str):
        return value.strip() == ""
    if isinstance(value, list | dict):
        return len(value) == 0
    return False


def _humanize_key(key: str) -> str:
    """``camelCase``/``snake_case`` → ``"Title Case"``.

    Mirrors ``humanizeKey`` in the viewer.
    """
    if not key:
        return ""
    s = key.replace("_", " ")
    out: list[str] = []
    for i, c in enumerate(s):
        if c.isupper() and i > 0 and (s[i - 1].islower() or s[i - 1].isdigit()):
            out.append(" ")
        out.append(c)
    cleaned = "".join(out).strip()
    return cleaned[:1].upper() + cleaned[1:] if cleaned else ""


def _primary_text(item: dict[str, Any]) -> tuple[str | None, str | None]:
    """Return (field_name, value) for the first non-empty primary-text field, else (None, None)."""
    for k in PRIMARY_TEXT_FIELDS:
        v = item.get(k)
        if isinstance(v, str) and v.strip():
            return k, v
    return None, None


def _visible_keys(item: dict[str, Any]) -> list[str]:
    """Item keys with hidden IDs filtered out, in insertion order."""
    return [k for k in item if k not in HIDDEN_FIELDS]


def _is_content_wrapper(value: Any) -> str | None:
    """If ``value`` is ``{"content": "..."}`` shape, return the inner string. Otherwise None.

    All four seed templates wrap ``executiveSummary`` like this; AI-generated
    content may extend the pattern.
    """
    if not isinstance(value, dict):
        return None
    visible = _visible_keys(value)
    if visible == ["content"] and isinstance(value["content"], str):
        return value["content"]
    return None


def _format_currency(raw: Any) -> str | None:
    """Format a number-like value as ``$1,234.56``. Returns None on failure or empty."""
    if raw is None or raw == "":
        return None
    try:
        return f"${float(raw):,.2f}"
    except (TypeError, ValueError):
        return str(raw)


# ── Table-promotion logic ─────────────────────────────────────────────────────


def _is_table_candidate(items: list[Any]) -> tuple[bool, list[str]]:
    """Decide whether a ``list[dict]`` should render as a Word table.

    Returns ``(is_candidate, column_keys)`` where ``column_keys`` is the
    ordered list of visible keys shared across all items.
    """
    if not items or len(items) < 2:
        return False, []
    if not all(isinstance(it, dict) for it in items):
        return False, []

    first_keys = _visible_keys(items[0])
    if not first_keys:
        return False, []
    if len(first_keys) > _MAX_TABLE_COLUMNS:
        return False, []
    if len(first_keys) < 2:
        # Items with only a primary-text column render better as bullets.
        return False, []

    key_set = set(first_keys)
    for it in items[1:]:
        if set(_visible_keys(it)) != key_set:
            return False, []

    # Guard against long cell text that wraps badly in narrow columns.
    for it in items:
        for k in first_keys:
            v = it.get(k)
            if isinstance(v, str) and len(v) > _MAX_TABLE_CELL_CHARS:
                return False, []

    return True, first_keys


def _is_simple_text_array(items: list[Any]) -> bool:
    """True when every item is a dict whose only visible field is the primary-text field.

    E.g., ``[{"id": "is-1", "text": "..."}, {"id": "is-2", "text": "..."}]`` —
    renders as a plain bullet list rather than as object cards.
    """
    if not items:
        return False
    for it in items:
        if not isinstance(it, dict):
            return False
        visible = _visible_keys(it)
        if len(visible) != 1:
            return False
        prim_key, prim_val = _primary_text(it)
        if prim_key != visible[0]:
            return False
        if not isinstance(prim_val, str):
            return False
    return True


# ── Word styling helpers ──────────────────────────────────────────────────────


def _set_default_styles(doc: Any) -> None:
    """Apply Calibri 11 / 1.15 line-spacing to the Normal style.

    Word users' personal defaults vary; pinning this here keeps the doc
    consistent across machines.
    """
    from docx.shared import Pt

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15


def _add_header_footer(doc: Any, title: str) -> None:
    """Add SoW title to the page header and an auto-updating page number to the footer."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = doc.sections[0]

    header_para = section.header.paragraphs[0]
    header_para.text = title or "Statement of Work"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_para.runs:
        run.italic = True

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _bold_cell(cell: Any) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def _render_table(doc: Any, items: list[dict[str, Any]], columns: list[str]) -> None:
    """Render a uniform list of dicts as a Word table with a bold header row."""
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, k in enumerate(columns):
        header_cells[i].text = _humanize_key(k)
        _bold_cell(header_cells[i])
    for item in items:
        row = table.add_row().cells
        for i, k in enumerate(columns):
            v = item.get(k, "")
            if isinstance(v, dict | list):
                row[i].text = json.dumps(v, ensure_ascii=False)
            elif isinstance(v, bool):
                row[i].text = "Yes" if v else "No"
            elif v is None:
                row[i].text = ""
            else:
                row[i].text = str(v)


# ── Risks special-case ────────────────────────────────────────────────────────


def _render_risk_item(doc: Any, risk: Any) -> None:
    """Render a single risk with optional severity tag + mitigation sub-line."""
    if isinstance(risk, str):
        doc.add_paragraph(risk, style="List Bullet")
        return
    if not isinstance(risk, dict):
        doc.add_paragraph(str(risk), style="List Bullet")
        return

    severity = risk.get("severity") or risk.get("level") or ""
    description = (
        risk.get("description") or risk.get("text") or risk.get("risk") or risk.get("title") or ""
    )
    mitigation = risk.get("mitigation") or ""

    if not description and not severity and not mitigation:
        # Fall back to JSON to avoid silently swallowing structured-but-unknown shapes.
        doc.add_paragraph(json.dumps(risk, ensure_ascii=False), style="List Bullet")
        return

    prefix = f"[{str(severity).upper()}] " if severity else ""
    doc.add_paragraph(f"{prefix}{description}".strip(), style="List Bullet")
    if mitigation:
        p = doc.add_paragraph()
        run = p.add_run(f"   Mitigation: {mitigation}")
        run.italic = True


def _render_risks_list(doc: Any, items: list[Any]) -> None:
    for r in items:
        _render_risk_item(doc, r)


# ── Recursive renderer ────────────────────────────────────────────────────────


def _render_card(doc: Any, item: dict[str, Any]) -> None:
    """Render a single dict item as a "card": bold primary text + sub-fields."""
    prim_key, prim_val = _primary_text(item)
    if prim_key and prim_val:
        p = doc.add_paragraph()
        p.add_run(prim_val).bold = True

    sub_entries = [
        (k, v)
        for k, v in item.items()
        if k != prim_key and k not in HIDDEN_FIELDS and not _is_empty(v)
    ]
    if not sub_entries:
        return

    # Render sub-fields as a small 2-col table for readability.
    sub_table = doc.add_table(rows=len(sub_entries), cols=2)
    sub_table.style = "Table Grid"
    for i, (k, v) in enumerate(sub_entries):
        sub_table.rows[i].cells[0].text = _humanize_key(k)
        _bold_cell(sub_table.rows[i].cells[0])
        if isinstance(v, dict | list):
            sub_table.rows[i].cells[1].text = json.dumps(v, ensure_ascii=False)
        elif isinstance(v, bool):
            sub_table.rows[i].cells[1].text = "Yes" if v else "No"
        elif v is None:
            sub_table.rows[i].cells[1].text = ""
        else:
            sub_table.rows[i].cells[1].text = str(v)
    # Spacer paragraph between cards.
    doc.add_paragraph()


def _render_value(
    doc: Any,
    value: Any,
    *,
    section_number: str,
    depth: int = 0,
    parent_key: str | None = None,
) -> None:
    """Append paragraphs/lists/tables for ``value`` into ``doc``.

    ``section_number`` is the dotted prefix used for sub-section headings
    (e.g. ``"2"``, ``"2.1"``).  ``parent_key`` is the JSON key whose value is
    being rendered — used to special-case risks.
    """
    if _is_empty(value):
        return

    # Unwrap ``{"content": "..."}`` envelopes.
    wrapped = _is_content_wrapper(value)
    if wrapped is not None:
        doc.add_paragraph(wrapped)
        return

    if isinstance(value, str):
        doc.add_paragraph(value)
        return

    if isinstance(value, bool):
        doc.add_paragraph("Yes" if value else "No")
        return

    if isinstance(value, int | float):
        doc.add_paragraph(str(value))
        return

    if isinstance(value, list):
        # Risks get their own special-case renderer (severity tag + mitigation).
        if parent_key == "risks":
            _render_risks_list(doc, value)
            return

        # All-string array → bullet list.
        if all(isinstance(v, str) for v in value):
            for s in value:
                doc.add_paragraph(s, style="List Bullet")
            return

        # All-dict array with only primary-text field → bullet list of that text.
        if _is_simple_text_array(value):
            for it in value:
                _, prim_val = _primary_text(it)
                doc.add_paragraph(prim_val or "", style="List Bullet")
            return

        # Uniform-shape table candidate.
        if all(isinstance(v, dict) for v in value):
            ok, cols = _is_table_candidate(value)
            if ok:
                _render_table(doc, value, cols)
                return
            # Heterogeneous or long-text → card layout.
            for it in value:
                _render_card(doc, it)
            return

        # Mixed types — render each through the recursive renderer.
        for it in value:
            _render_value(doc, it, section_number=section_number, depth=depth + 1)
        return

    if isinstance(value, dict):
        entries = [(k, v) for k, v in value.items() if k not in HIDDEN_FIELDS and not _is_empty(v)]
        if not entries:
            return
        for i, (k, v) in enumerate(entries, start=1):
            sub_number = f"{section_number}.{i}"
            label = SECTION_LABELS.get(k) or _humanize_key(k)
            # Skip the sub-heading when this is a single ``{"content": "..."}``
            # envelope (already unwrapped above), or when the child wraps one.
            wrapped_child = _is_content_wrapper(v)
            if wrapped_child is not None:
                heading_level = min(2 + depth, 4)
                doc.add_heading(f"{sub_number} {label}", level=heading_level)
                doc.add_paragraph(wrapped_child)
                continue
            heading_level = min(2 + depth, 4)
            doc.add_heading(f"{sub_number} {label}", level=heading_level)
            _render_value(
                doc,
                v,
                section_number=sub_number,
                depth=depth + 1,
                parent_key=k,
            )
        return

    # Fallback for unknown types (datetime, Decimal, custom objects).
    doc.add_paragraph(str(value))


# ── Cover page ────────────────────────────────────────────────────────────────


def _cover_metadata_rows(sow: dict[str, Any], content: dict[str, Any]) -> list[tuple[str, str]]:
    """Build the (label, value) pairs for the cover-page metadata table.

    Prefers values from ``content`` (matching the viewer's summary card)
    before falling back to row columns.
    """

    def from_content_or_sow(content_key: str, *sow_keys: str) -> Any:
        v = content.get(content_key)
        if not _is_empty(v):
            return v
        for sk in sow_keys:
            sv = sow.get(sk)
            if not _is_empty(sv):
                return sv
        return None

    title = from_content_or_sow("sowTitle", "title")
    customer = from_content_or_sow("customerName", "customer_name")
    methodology = from_content_or_sow("deliveryMethodology", "methodology")
    opportunity_id = from_content_or_sow("opportunityId", "opportunity_id")
    deal_value_raw = from_content_or_sow("dealValue", "deal_value")
    status = from_content_or_sow("status", "status")
    deal_value = _format_currency(deal_value_raw)

    rows: list[tuple[str, str]] = [
        ("Title", str(title or "")),
        ("Customer", str(customer or "")),
        ("Methodology", str(methodology or "")),
        ("Opportunity ID", str(opportunity_id or "")),
    ]
    if deal_value is not None:
        rows.append(("Deal Value", deal_value))
    if status:
        rows.append(("Status", str(status)))
    rows.append(("Generated", datetime.now(UTC).strftime("%B %d, %Y")))
    return rows


def _render_cover(doc: Any, sow: dict[str, Any], content: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title_para = doc.add_heading("STATEMENT OF WORK", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows = _cover_metadata_rows(sow, content)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        _bold_cell(table.rows[i].cells[0])
        table.rows[i].cells[1].text = value

    doc.add_paragraph()
    doc.add_page_break()


# ── Section ordering ──────────────────────────────────────────────────────────


def _ordered_section_keys(content: dict[str, Any]) -> list[str]:
    """Return non-empty, non-metadata content keys in canonical reading order.

    Unknown keys are appended at the end so unfamiliar templates aren't
    silently truncated.
    """
    metadata = set(METADATA_FIELDS)
    present = [k for k in content if k not in metadata and not _is_empty(content[k])]
    seen: set[str] = set()
    ordered: list[str] = []
    for k in SECTION_ORDER:
        if k in present:
            ordered.append(k)
            seen.add(k)
    for k in present:
        if k not in seen:
            ordered.append(k)
    return ordered


def _nested_keys(value: Any) -> set[str]:
    """Top-level child keys of a dict that also appear in ``SECTION_ORDER``."""
    if not isinstance(value, dict):
        return set()
    return {k for k in value if k in SECTION_ORDER and not _is_empty(value[k])}


def _render_sections(doc: Any, content: dict[str, Any]) -> None:
    keys = _ordered_section_keys(content)
    rendered_via_nesting: set[str] = set()
    section_index = 0
    for k in keys:
        if k in rendered_via_nesting:
            continue
        section_index += 1
        label = SECTION_LABELS.get(k) or _humanize_key(k)
        doc.add_heading(f"{section_index}. {label.upper()}", level=1)

        value = content[k]
        # Record any child keys that this section will absorb so we don't
        # render them again at the top level.
        rendered_via_nesting.update(_nested_keys(value))

        _render_value(
            doc,
            value,
            section_number=str(section_index),
            depth=0,
            parent_key=k,
        )


# ── Appendices (unchanged from the previous implementation) ───────────────────


def _safe_json(value: Any) -> Any:
    """Tolerant of values that are either already-parsed or still JSON strings."""
    if value is None or isinstance(value, list | dict):
        return value
    if isinstance(value, str):
        try:
            return json.loads(value)
        except (ValueError, TypeError):
            return value
    return value


def _render_appendix_a(doc: Any, review_results: list[dict[str, Any]]) -> None:
    conditions: list[str] = []
    for r in review_results:
        cond = _safe_json(r.get("conditions"))
        if isinstance(cond, list):
            conditions.extend(str(c) for c in cond if c)
        elif isinstance(cond, str) and cond:
            conditions.append(cond)
    if not conditions:
        return
    doc.add_page_break()
    doc.add_heading("APPENDIX A: REVIEW CONDITIONS", level=1)
    for c in conditions:
        doc.add_paragraph(c, style="List Bullet")


def _render_appendix_b(doc: Any, review_results: list[dict[str, Any]]) -> None:
    if not review_results:
        return
    doc.add_page_break()
    doc.add_heading("APPENDIX B: APPROVAL CHAIN", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(("Reviewer", "Role", "Decision", "Date")):
        table.rows[0].cells[i].text = h
        _bold_cell(table.rows[0].cells[i])
    for r in review_results:
        cells = table.add_row().cells
        cells[0].text = str(r.get("reviewer") or "")
        cells[1].text = str(r.get("reviewer_role") or r.get("review_stage") or "")
        cells[2].text = str(r.get("decision") or "")
        date_val = r.get("reviewed_at")
        if isinstance(date_val, datetime):
            cells[3].text = date_val.strftime("%b %d, %Y")
        elif date_val:
            cells[3].text = str(date_val)[:10]
        else:
            cells[3].text = ""


# ── Public entry point ───────────────────────────────────────────────────────


def render_sow_to_docx(
    sow: dict[str, Any],
    content: dict[str, Any],
    review_results: list[dict[str, Any]] | None = None,
) -> bytes:
    """Build a DOCX document mirroring the in-app SoW viewer. Returns raw bytes."""
    try:
        from docx import Document
    except ImportError as err:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx",
        ) from err

    doc = Document()
    _set_default_styles(doc)
    title_for_header = content.get("sowTitle") or sow.get("title") or "Statement of Work"
    _add_header_footer(doc, str(title_for_header))

    _render_cover(doc, sow, content or {})
    _render_sections(doc, content or {})

    review_results = review_results or []
    _render_appendix_a(doc, review_results)
    _render_appendix_b(doc, review_results)

    # TODO: optional auto-updating TOC via <w:fldSimple instr="TOC \o '1-3' \h \z \u"/>.
    # Skipped in v1 because LibreOffice often renders the field empty and Word
    # requires a manual F9 refresh on first open; Navigation Pane covers the
    # in-app reading use-case.

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
